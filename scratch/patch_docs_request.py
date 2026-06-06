import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def patch_volume_iii(doc_path):
    if not os.path.exists(doc_path):
        print(f"Skipping (not found): {doc_path}")
        return False
        
    doc = Document(doc_path)
    
    # 1. Append the Action Engine explanation after the Action Engine/Mitigation sections
    # We will search for a paragraph with text "Cross-Platform Resource Mitigation Strategy" or similar, or just append it at the end of Section 3.
    # To be safe and ensure the structure is correct, we can find paragraphs containing "Mitigation Strategy" and append after them,
    # or append it before the "Elite User Experience" section.
    
    # Let's add the Action Engine paragraph first.
    p_ae_head = doc.add_paragraph()
    p_ae_head.style = 'Heading 2'
    r_ae_head = p_ae_head.add_run("Kernel Thread Protection & Thread-Safe UI Interception")
    r_ae_head.bold = True
    r_ae_head.font.name = "Arial"
    r_ae_head.font.size = Pt(13)
    
    p_ae_body = doc.add_paragraph()
    p_ae_body.style = 'Normal'
    r_ae_body = p_ae_body.add_run(
        "The SRO prevents UI freezes by intercepting runaway processes before they can starve the Windows kernel's "
        "interrupt handlers and compositor threads. When the GRU model predicts a bottleneck with \u226580% confidence, "
        "the Action Engine either:\n"
        "  - Suspends the process entirely (SIGSTOP / NtSuspendThread), or\n"
        "  - On supported systems, isolates the heavy task to efficiency cores (macOS) or removes its access to Core 0/1 (Windows).\n"
        "This ensures that critical UI threads (explorer.exe, dwm.exe, csrss.exe) always receive CPU time, eliminating the 'Not Responding' experience."
    )
    r_ae_body.font.name = "Arial"
    r_ae_body.font.size = Pt(11)

    # 2. Add "Elite User Experience & Help System" subsection
    p_ux_head = doc.add_paragraph()
    p_ux_head.style = 'Heading 2'
    r_ux_head = p_ux_head.add_run("Elite User Experience & Help System")
    r_ux_head.bold = True
    r_ux_head.font.name = "Arial"
    r_ux_head.font.size = Pt(13)
    
    p_ux_intro = doc.add_paragraph()
    p_ux_intro.style = 'Normal'
    r_ux_intro = p_ux_intro.add_run(
        "To ensure high user accessibility and low cognitive friction, the system implements an multi-tiered, in-app "
        "documentation framework described below:"
    )
    r_ux_intro.font.name = "Arial"
    r_ux_intro.font.size = Pt(11)
    
    # Add table: In-App Help & Documentation Strategy
    table_data = [
        ("Tooltips", "Hover over any UI element (\u24d8 icon)", "1-sentence explanation"),
        ("Quick Start Guide (1 page PDF)", "In the installer folder or Help menu", "5 steps: Install \u2192 Launch \u2192 Wait 60s \u2192 Toggle Auto-Pilot \u2192 Done"),
        ("In-app \"Help\" sidebar", "Click \u2754 icon, slides out from right", "FAQ: Why suspend not kill? What is whitelist? How to restore?"),
        ("Video demo link", "Help menu \u2192 \"Watch 1-min demo\"", "YouTube unlisted or local MP4 (show live responsiveness)")
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ["Format", "Where", "Content"]
    for i, h_text in enumerate(headers):
        hdr_cells[i].text = h_text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(11)
                
    for fmt, whr, cnt in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = fmt
        row_cells[1].text = whr
        row_cells[2].text = cnt
        for i in range(3):
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    
    # Add spacing after table
    p_space = doc.add_paragraph()
    p_space.style = 'Normal'
    
    p_ux_conclusion = doc.add_paragraph()
    p_ux_conclusion.style = 'Normal'
    r_ux_conclusion = p_ux_conclusion.add_run(
        "Do NOT force the user to read a 20-page document. That is the opposite of an \u2018elite\u2019 user experience."
    )
    r_ux_conclusion.italic = True
    r_ux_conclusion.bold = True
    r_ux_conclusion.font.name = "Arial"
    r_ux_conclusion.font.size = Pt(11)
    
    doc.save(doc_path)
    print(f"Successfully patched {os.path.basename(doc_path)} with new SRO mitigation and Elite UX tables.")
    return True

def main():
    docs_dir = "/Users/user/Desktop/Final_year/docs"
    
    vol_iii_path = os.path.join(docs_dir, "Volume_III_Pipeline_UI_and_Verification.docx")
    master_path = os.path.join(docs_dir, "System_Resource_Optimizer_Documentation.docx")
    
    patch_volume_iii(vol_iii_path)
    patch_volume_iii(master_path)

if __name__ == "__main__":
    main()
