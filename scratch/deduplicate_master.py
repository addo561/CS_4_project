import os
from docx import Document

def audit_document():
    master_path = "/Users/user/Desktop/Final_year/docs/System_Resource_Optimizer_Documentation.docx"
    if not os.path.exists(master_path):
        print(f"Error: Master document not found at {master_path}")
        return
        
    doc = Document(master_path)
    print("==================================================")
    print("      MASTER COMPENDIUM AUDIT & VERIFICATION     ")
    print("==================================================")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    
    # 1. Check title page occurrences
    univ_text = "KWAME NKRUMAH UNIVERSITY OF SCIENCE AND TECHNOLOGY"
    univ_count = 0
    student_text = "Lamptey Kwaku Abednego"
    student_count = 0
    
    for p in doc.paragraphs:
        if univ_text in p.text:
            univ_count += 1
        if student_text in p.text:
            student_count += 1
            
    print(f"University Title occurrences: {univ_count} (Expected: 1)")
    print(f"Student Listing occurrences: {student_count} (Expected: 1)")
    
    # 2. Check chapter heading sequences
    headings_found = []
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            headings_found.append(p.text)
            
    print("\nChapter Heading Structure:")
    for h in headings_found:
        print(f" - {h}")
        
    # 3. Check for any remaining raw LaTeX formulas
    latex_matches = []
    for idx, p in enumerate(doc.paragraphs):
        # Look for raw $ inline or $$ blocks
        text = p.text
        if "$$" in text:
            latex_matches.append((idx + 1, "Block", text[:60]))
        elif "$" in text and not re_math_ok(text):
            latex_matches.append((idx + 1, "Inline", text[:60]))
            
    print(f"\nRemaining Raw LaTeX Math Blocks: {len(latex_matches)}")
    for line, mtype, snippet in latex_matches[:5]:
        print(f" - Line {line} ({mtype}): '{snippet}...'")
        
    # 4. Check for scattered references
    references_headings = [h for h in headings_found if "references" in h.lower() or "bibliography" in h.lower()]
    print(f"\nReferences/Bibliography Headings: {len(references_headings)} (Expected: 1)")
    for rh in references_headings:
        print(f" - {rh}")
        
    print("==================================================")

def re_math_ok(text):
    # Some dollar signs might be standard text or already formatted runs
    # If the text has $ followed immediately by letters/symbols (like $x_t$), it's math.
    # Standard usage like $38.00 or $10 is financial/non-math.
    import re
    math_patterns = [r'\$[a-zA-Z]_[a-zA-Z0-9\+\-]', r'\$\\text', r'\$\\mathcal', r'\$\\sigma', r'\$\\tau']
    for pat in math_patterns:
        if re.search(pat, text):
            return False
    return True

if __name__ == "__main__":
    audit_document()
