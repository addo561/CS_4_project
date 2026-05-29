import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_math_in_paragraph(p):
    text = p.text
    # Quick filter to avoid processing regular text
    if "$" not in text and "\\text" not in text and "\\sigma" not in text and "\\mathcal" not in text and "\\tau" not in text and "\\ge" not in text:
        return False
        
    # Check if this is a block equation centered with $$
    if text.strip().startswith("$$") and text.strip().endswith("$$"):
        eq = text.strip().strip("$").strip()
        p.text = "" # Clear runs
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 1. Temperature Simulation Target Formula
        if "T_{\\text{target}}" in eq:
            r1 = p.add_run("T")
            r1.italic = True
            r2 = p.add_run("target")
            r2.font.subscript = True
            p.add_run(" = 38.0 + (")
            r3 = p.add_run("U")
            r3.italic = True
            r4 = p.add_run("cpu")
            r4.font.subscript = True
            p.add_run(" × 0.45) + (")
            r5 = p.add_run("U")
            r5.italic = True
            r6 = p.add_run("mem")
            r6.font.subscript = True
            p.add_run(" × 0.1)")
            return True
            
        # 2. Temperature Simulation Lag Decay Formula
        elif "T_{t+1}" in eq:
            r1 = p.add_run("T")
            r1.italic = True
            r2 = p.add_run("t+1")
            r2.font.subscript = True
            p.add_run(" = ")
            r3 = p.add_run("T")
            r3.italic = True
            r4 = p.add_run("t")
            r4.font.subscript = True
            p.add_run(" + (")
            r5 = p.add_run("T")
            r5.italic = True
            r6 = p.add_run("target")
            r6.font.subscript = True
            p.add_run(" - ")
            r7 = p.add_run("T")
            r7.italic = True
            r8 = p.add_run("t")
            r8.font.subscript = True
            p.add_run(") × 0.1")
            return True
            
        # 3. GRU Update Gate Formula
        elif "z_t = \\sigma" in eq or "z_t=\\sigma" in eq:
            r1 = p.add_run("z")
            r1.italic = True
            r2 = p.add_run("t")
            r2.font.subscript = True
            p.add_run(" = σ(")
            r3 = p.add_run("W")
            r3.italic = True
            r4 = p.add_run("z")
            r4.font.subscript = True
            p.add_run(" · [")
            r5 = p.add_run("h")
            r5.italic = True
            r6 = p.add_run("t-1")
            r6.font.subscript = True
            p.add_run(", ")
            r7 = p.add_run("x")
            r7.italic = True
            r8 = p.add_run("t")
            r8.font.subscript = True
            p.add_run("] + ")
            r9 = p.add_run("b")
            r9.italic = True
            r10 = p.add_run("z")
            r10.font.subscript = True
            p.add_run(")")
            return True
            
        # 4. GRU Reset Gate Formula
        elif "r_t = \\sigma" in eq or "r_t=\\sigma" in eq:
            r1 = p.add_run("r")
            r1.italic = True
            r2 = p.add_run("t")
            r2.font.subscript = True
            p.add_run(" = σ(")
            r3 = p.add_run("W")
            r3.italic = True
            r4 = p.add_run("r")
            r4.font.subscript = True
            p.add_run(" · [")
            r5 = p.add_run("h")
            r5.italic = True
            r6 = p.add_run("t-1")
            r6.font.subscript = True
            p.add_run(", ")
            r7 = p.add_run("x")
            r7.italic = True
            r8 = p.add_run("t")
            r8.font.subscript = True
            p.add_run("] + ")
            r9 = p.add_run("b")
            r9.italic = True
            r10 = p.add_run("r")
            r10.font.subscript = True
            p.add_run(")")
            return True
            
        # 5. GRU Candidate Hidden State Formula
        elif "\\tilde{h}_t = \\tanh" in eq or "\\tilde{h}_t=\\tanh" in eq:
            r1 = p.add_run("h̃")
            r1.italic = True
            r2 = p.add_run("t")
            r2.font.subscript = True
            p.add_run(" = tanh(")
            r3 = p.add_run("W")
            r3.italic = True
            r4 = p.add_run("h")
            r4.font.subscript = True
            p.add_run(" · [")
            r5 = p.add_run("r")
            r5.italic = True
            r6 = p.add_run("t")
            r6.font.subscript = True
            p.add_run(" ⊙ ")
            r7 = p.add_run("h")
            r7.italic = True
            r8 = p.add_run("t-1")
            r8.font.subscript = True
            p.add_run(", ")
            r9 = p.add_run("x")
            r9.italic = True
            r10 = p.add_run("t")
            r10.font.subscript = True
            p.add_run("] + ")
            r11 = p.add_run("b")
            r11.italic = True
            r12 = p.add_run("h")
            r12.font.subscript = True
            p.add_run(")")
            return True
            
        # 6. GRU Final Hidden State Formula
        elif "h_t = (1 - z_t)" in eq or "h_t=(1-z_t)" in eq:
            r1 = p.add_run("h")
            r1.italic = True
            r2 = p.add_run("t")
            r2.font.subscript = True
            p.add_run(" = (1 - ")
            r3 = p.add_run("z")
            r3.italic = True
            r4 = p.add_run("t")
            r4.font.subscript = True
            p.add_run(") ⊙ ")
            r5 = p.add_run("h")
            r5.italic = True
            r6 = p.add_run("t-1")
            r6.font.subscript = True
            p.add_run(" + ")
            r7 = p.add_run("z")
            r7.italic = True
            r8 = p.add_run("t")
            r8.font.subscript = True
            p.add_run(" ⊙ ")
            r9 = p.add_run("h̃")
            r9.italic = True
            r10 = p.add_run("t")
            r10.font.subscript = True
            return True
            
        # 7. Composite Loss Formula
        elif "\\mathcal{L}_{\\text{total}}" in eq:
            r1 = p.add_run("L")
            r1.italic = True
            r2 = p.add_run("total")
            r2.font.subscript = True
            p.add_run(" = ")
            r3 = p.add_run("L")
            r3.italic = True
            r4 = p.add_run("MSE")
            r4.font.subscript = True
            p.add_run("(reg_out, ")
            r5 = p.add_run("y")
            r5.italic = True
            r6 = p.add_run("reg")
            r6.font.subscript = True
            p.add_run(") + ")
            r7 = p.add_run("L")
            r7.italic = True
            r8 = p.add_run("BCE")
            r8.font.subscript = True
            p.add_run("(clf_out, ")
            r9 = p.add_run("y")
            r9.italic = True
            r10 = p.add_run("clf")
            r10.font.subscript = True
            p.add_run(")")
            return True
            
    # For inline math expressions (containing single $ markers)
    if "$" in text:
        parts = text.split("$")
        p.text = "" # Clear runs
        for i, part in enumerate(parts):
            if i % 2 == 0:
                p.add_run(part)
            else:
                math = part.strip()
                if not math:
                    continue
                # Map standard inline LaTeX variables to beautiful styled runs
                if math == "P(B_{t+1}) > \\tau" or "P(B_{t+1})" in math:
                    r1 = p.add_run("P")
                    r1.italic = True
                    p.add_run("(")
                    r2 = p.add_run("B")
                    r2.italic = True
                    r3 = p.add_run("t+1")
                    r3.font.subscript = True
                    p.add_run(") > τ")
                elif math == "x_1, x_2, ..., x_t":
                    r1 = p.add_run("x")
                    r1.italic = True
                    r2 = p.add_run("1")
                    r2.font.subscript = True
                    p.add_run(", ")
                    r3 = p.add_run("x")
                    r3.italic = True
                    r4 = p.add_run("2")
                    r4.font.subscript = True
                    p.add_run(", ..., ")
                    r5 = p.add_run("x")
                    r5.italic = True
                    r6 = p.add_run("t")
                    r6.font.subscript = True
                elif "x_t \\in \\mathbb{R}^F" in math or "x_t" in math and "\\mathbb{R}^F" in math:
                    r1 = p.add_run("x")
                    r1.italic = True
                    r2 = p.add_run("t")
                    r2.font.subscript = True
                    p.add_run(" ∈ ℝ")
                    r3 = p.add_run("F")
                    r3.font.superscript = True
                elif "h_t \\in \\mathbb{R}^H" in math or "h_t" in math and "\\mathbb{R}^H" in math:
                    r1 = p.add_run("h")
                    r1.italic = True
                    r2 = p.add_run("t")
                    r2.font.subscript = True
                    p.add_run(" ∈ ℝ")
                    r3 = p.add_run("H")
                    r3.font.superscript = True
                elif math == "z_t":
                    r1 = p.add_run("z")
                    r1.italic = True
                    r2 = p.add_run("t")
                    r2.font.subscript = True
                elif math == "r_t":
                    r1 = p.add_run("r")
                    r1.italic = True
                    r2 = p.add_run("t")
                    r2.font.subscript = True
                elif math == "h_t":
                    r1 = p.add_run("h")
                    r1.italic = True
                    r2 = p.add_run("t")
                    r2.font.subscript = True
                elif math == "h_{t-1}":
                    r1 = p.add_run("h")
                    r1.italic = True
                    r2 = p.add_run("t-1")
                    r2.font.subscript = True
                elif math == "\\tilde{h}_t" or math == "\\tilde{h}_t":
                    r1 = p.add_run("h̃")
                    r1.italic = True
                    r2 = p.add_run("t")
                    r2.font.subscript = True
                elif math == "W_z":
                    r1 = p.add_run("W")
                    r1.italic = True
                    r2 = p.add_run("z")
                    r2.font.subscript = True
                elif math == "b_z":
                    r1 = p.add_run("b")
                    r1.italic = True
                    r2 = p.add_run("z")
                    r2.font.subscript = True
                elif math == "W_r":
                    r1 = p.add_run("W")
                    r1.italic = True
                    r2 = p.add_run("r")
                    r2.font.subscript = True
                elif math == "b_r":
                    r1 = p.add_run("b")
                    r1.italic = True
                    r2 = p.add_run("r")
                    r2.font.subscript = True
                elif math == "W_h":
                    r1 = p.add_run("W")
                    r1.italic = True
                    r2 = p.add_run("h")
                    r2.font.subscript = True
                elif math == "b_h":
                    r1 = p.add_run("b")
                    r1.italic = True
                    r2 = p.add_run("h")
                    r2.font.subscript = True
                elif math == "10^{-3}":
                    p.add_run("10")
                    r = p.add_run("-3")
                    r.font.superscript = True
                elif math in ("[0, 1]", "[0,1]"):
                    p.add_run("[0, 1]")
                elif math in ("[-1, 1]", "[-1,1]"):
                    p.add_run("[-1, 1]")
                elif math == "\\mathcal{L}_{\\text{MSE}}":
                    r1 = p.add_run("L")
                    r1.italic = True
                    r2 = p.add_run("MSE")
                    r2.font.subscript = True
                elif math == "\\mathcal{L}_{\\text{BCE}}":
                    r1 = p.add_run("L")
                    r1.italic = True
                    r2 = p.add_run("BCE")
                    r2.font.subscript = True
                elif math == "\\mathcal{L}_{\\text{total}}":
                    r1 = p.add_run("L")
                    r1.italic = True
                    r2 = p.add_run("total")
                    r2.font.subscript = True
                elif math in ("y_{\\text{reg}}", "y_{reg}"):
                    r1 = p.add_run("y")
                    r1.italic = True
                    r2 = p.add_run("reg")
                    r2.font.subscript = True
                elif math in ("y_{\\text{clf}}", "y_{clf}"):
                    r1 = p.add_run("y")
                    r1.italic = True
                    r2 = p.add_run("clf")
                    r2.font.subscript = True
                elif math == "t+30":
                    r1 = p.add_run("t")
                    r1.italic = True
                    p.add_run("+30")
                elif math == "t+1":
                    r1 = p.add_run("t")
                    r1.italic = True
                    p.add_run("+1")
                elif math == "t+H":
                    r1 = p.add_run("t")
                    r1.italic = True
                    p.add_run("+")
                    r2 = p.add_run("H")
                    r2.italic = True
                elif math in ("\\ge 80\\%", "\\ge 80%"):
                    p.add_run("≥ 80%")
                elif math in ("\\ge 65\\%", "\\ge 65%"):
                    p.add_run("≥ 65%")
                elif math in ("\\ge 85\\%", "\\ge 85%"):
                    p.add_run("≥ 85%")
                elif math in ("\\ge 55\\%", "\\ge 55%"):
                    p.add_run("≥ 55%")
                else:
                    # Clean simple variable styling
                    r = p.add_run(math)
                    r.italic = True
        return True
    return False

def clean_file(path):
    print(f"Parsing and fixing equations in: {os.path.basename(path)}")
    doc = Document(path)
    fixed_count = 0
    for p in doc.paragraphs:
        if fix_math_in_paragraph(p):
            fixed_count += 1
    if fixed_count > 0:
        doc.save(path)
        print(f"Successfully saved with {fixed_count} equations beautifully formatted!")
    else:
        print("No equations needed fixing.")

if __name__ == "__main__":
    docs_dir = "/Users/user/Desktop/Final_year/docs"
    # Fix all modular docx files and the merged compendium!
    files = [
        "introductions.docx",
        "dataset_collection.docx",
        "model_documentation.docx",
        "data_pipeline.docx",
        "ui_documentation.docx",
        "mitigation_strategy.docx",
        "presentation_deliverables.docx",
        "System_Resource_Optimizer_Documentation.docx"
    ]
    for filename in files:
        full_path = os.path.join(docs_dir, filename)
        if os.path.exists(full_path):
            clean_file(full_path)
