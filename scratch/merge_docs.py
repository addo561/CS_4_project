import os
import sys
import re
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.table import Table

def clean_ref_text(text):
    # Remove leading brackets and numbers, e.g. [1], [1] , [ 1 ]
    cleaned = re.sub(r'^\[\s*\d+\s*\]\s*', '', text)
    cleaned = cleaned.strip()
    return cleaned

def get_chapter_heading(chapter_num, original_text):
    # Map each chapter to a highly professional, academic chapter title
    titles = {
        1: "Chapter 1: Project Introduction & High-Level System Architecture",
        2: "Chapter 2: Telemetry Dataset Collection & Preprocessing",
        3: "Chapter 3: Gated Recurrent Unit (GRU) Model Design & Training",
        4: "Chapter 4: Real-Time Data Pipeline & Action Engine",
        5: "Chapter 5: User Interface Design & Asynchronous Architecture",
        6: "Chapter 6: Cross-Platform Resource Mitigation Strategy",
        7: "Chapter 7: System Verification, Live Demonstration & Presentation Deliverables"
    }
    return titles.get(chapter_num, f"Chapter {chapter_num}: {original_text}")

def create_cover_page(doc):
    print("Generating elegant master cover page...")
    
    # 1. University Header
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_univ = p_univ.add_run("KWAME NKRUMAH UNIVERSITY OF SCIENCE AND TECHNOLOGY\n")
    r_univ.bold = True
    r_univ.font.size = Pt(14)
    r_univ.font.name = "Arial"
    
    p_fac = doc.add_paragraph()
    p_fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_fac = p_fac.add_run("FACULTY OF PHYSICAL AND COMPUTATIONAL SCIENCES\nDEPARTMENT OF COMPUTER SCIENCE\n")
    r_fac.bold = True
    r_fac.font.size = Pt(11)
    r_fac.font.name = "Arial"
    
    # Spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # 2. Main Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("LIGHTWEIGHT AI-POWERED SYSTEM\nRESOURCE OPTIMIZER\n")
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.name = "Arial"
    r_title.font.color.rgb = RGBColor(13, 17, 23) # Sleek deep dark color
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Final Year Project Compendium & Unified Dissertation\n")
    r_sub.italic = True
    r_sub.font.size = Pt(12)
    r_sub.font.name = "Arial"
    
    # Spacing
    for _ in range(5):
        doc.add_paragraph()
        
    # 3. Prepared By
    p_prep = doc.add_paragraph()
    p_prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_prep = p_prep.add_run("Prepared by:\n")
    r_prep.bold = True
    r_prep.font.size = Pt(11)
    r_prep.font.name = "Arial"
    
    p_students = doc.add_paragraph()
    p_students.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_std = p_students.add_run(
        "Lamptey Kwaku Abednego — 3398122\n"
        "Tugbah Lily Ama Mawuena — 3416522\n"
        "Korli Larry Addo — 3395922\n"
    )
    r_std.font.size = Pt(11)
    r_std.font.name = "Arial"
    
    # Spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # 4. Footer Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run("Date: May 2026")
    r_date.bold = True
    r_date.font.size = Pt(11)
    r_date.font.name = "Arial"

def merge_docx_files(source_files, output_path):
    print(f"Starting non-redundant merge of {len(source_files)} files...")
    
    if not source_files:
        print("No files to merge.")
        return False
        
    # Start with a base document from the first source file to inherit styles, margins, page setups
    merged_doc = Document(source_files[0])
    
    # Clear the entire body of the base document to start fresh
    print("Clearing base template body while preserving styles...")
    body = merged_doc.element.body
    for child in list(body):
        if not child.tag.endswith('sectPr'):
            body.remove(child)
            
    # Generate the cover page
    create_cover_page(merged_doc)
    
    unique_ref_texts = set()
    unique_ref_list = []
    chapter_num = 1
    
    for i, file_path in enumerate(source_files, start=1):
        print(f"Processing and merging file {i}: {os.path.basename(file_path)}")
        
        # Add a page break before each new chapter
        merged_doc.add_page_break()
        
        # Open source document
        src_doc = Document(file_path)
        
        reached_content = False
        in_references = False
        
        for child in list(src_doc.element.body):
            # Skip page-setup elements (sectPr)
            if child.tag.endswith('sectPr'):
                continue
                
            if child.tag.endswith('p'):
                p = Paragraph(child, src_doc)
                
                # Check for start of content (Heading 1)
                if not reached_content:
                    if p.style.name == 'Heading 1':
                        reached_content = True
                        new_heading = get_chapter_heading(chapter_num, p.text)
                        p.text = new_heading
                        print(f" -> Substituted Heading: '{new_heading}'")
                        chapter_num += 1
                    else:
                        # Skip title block paragraphs at the beginning
                        continue
                        
                # Check for references section
                if p.style.name == 'Heading 1' and 'references' in p.text.lower():
                    in_references = True
                    continue
                    
                if in_references:
                    ref_text = p.text.strip()
                    if ref_text:
                        cleaned = clean_ref_text(ref_text)
                        if cleaned and cleaned not in unique_ref_texts:
                            unique_ref_texts.add(cleaned)
                            unique_ref_list.append(cleaned)
                    continue
                    
            elif child.tag.endswith('tbl'):
                if not reached_content:
                    # Skip any tables before the first heading
                    continue
                if in_references:
                    # Skip tables in references
                    continue
            
            # Insert element before sectPr to maintain valid order
            sectPr = body.sectPr
            if sectPr is not None:
                idx = body.index(sectPr)
                body.insert(idx, child)
            else:
                body.append(child)
            
    # Add consolidated references chapter at the very end
    if unique_ref_list:
        print(f"Merging and sorting {len(unique_ref_list)} consolidated bibliography references...")
        merged_doc.add_page_break()
        
        p_ref_head = merged_doc.add_paragraph()
        p_ref_head.style = 'Heading 1'
        p_ref_head.text = "Chapter 8: Consolidated Bibliography & References"
        
        sorted_refs = sorted(unique_ref_list, key=lambda s: s.lower())
        
        for idx, ref_text in enumerate(sorted_refs, start=1):
            p_ref = merged_doc.add_paragraph()
            p_ref.style = 'Normal'
            p_ref.paragraph_format.left_indent = Pt(18) # Hanging indent feel
            p_ref.paragraph_format.first_line_indent = Pt(-18)
            
            # Bold citation index
            r_idx = p_ref.add_run(f"[{idx}] ")
            r_idx.bold = True
            r_idx.font.name = "Arial"
            
            r_text = p_ref.add_run(ref_text)
            r_text.font.name = "Arial"
            
    print(f"Saving merged document to: {output_path}")
    merged_doc.save(output_path)
    print("Merge completed successfully!")
    return True

if __name__ == "__main__":
    docs_dir = "/Users/user/Desktop/Final_year/docs"
    
    ordered_files = [
        os.path.join(docs_dir, "introductions.docx"),
        os.path.join(docs_dir, "dataset_collection.docx"),
        os.path.join(docs_dir, "model_documentation.docx"),
        os.path.join(docs_dir, "data_pipeline.docx"),
        os.path.join(docs_dir, "ui_documentation.docx"),
        os.path.join(docs_dir, "mitigation_strategy.docx"),
        os.path.join(docs_dir, "presentation_deliverables.docx")
    ]
    
    missing_files = [f for f in ordered_files if not os.path.exists(f)]
    if missing_files:
        print(f"Error: The following source files are missing: {missing_files}", file=sys.stderr)
        sys.exit(1)
        
    output_docx = os.path.join(docs_dir, "System_Resource_Optimizer_Documentation.docx")
    
    success = merge_docx_files(ordered_files, output_docx)
    if success:
        print(f"Merged document successfully created at: {output_docx}")
    else:
        sys.exit(1)
