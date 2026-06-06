import os
from docx import Document
from docx.shared import Pt

def add_custom_section(doc_path, heading_text, paragraphs):
    if not os.path.exists(doc_path):
        print(f"Skipping (not found): {doc_path}")
        return False
        
    doc = Document(doc_path)
    
    # Add a paragraph for heading
    p_head = doc.add_paragraph()
    p_head.style = 'Heading 2'
    r_head = p_head.add_run(heading_text)
    r_head.bold = True
    r_head.font.name = "Arial"
    r_head.font.size = Pt(13)
    
    # Add paragraphs
    for text in paragraphs:
        p = doc.add_paragraph()
        p.style = 'Normal'
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        
    doc.save(doc_path)
    print(f"Successfully patched {os.path.basename(doc_path)} with section '{heading_text}'")
    return True

def main():
    docs_dir = "/Users/user/Desktop/Final_year/docs"
    
    # 1. Patch Volume III
    vol_iii_path = os.path.join(docs_dir, "Volume_III_Pipeline_UI_and_Verification.docx")
    vol_iii_head = "macOS Launch Services Cache & Command-Line Notification Limitations"
    vol_iii_paras = [
        "During distribution and execution of compiled Flet GUI bundles on macOS, the operating system caches application icons based on the bundle identifier in its system Launch Services database. If a version of the application was previously registered with the default Flet icon, macOS will cache it, which may lead to the coexistence of both the custom application icon and the generic Flet icon in the taskbar/Dock. Users can resolve this issue by moving the built bundle out of the build/dist directory, renaming it, or resetting the macOS Dock/Finder services (via the command: killall Dock && killall Finder) to force the OS to rebuild its Launch Services cache and display the custom logo exclusively.",
        "Additionally, standard Python-based desktop notifications on macOS are sent using command-line AppleScript execution (via the osascript utility). Since the shell delegates execution to the system-wide 'osascript' command-line interpreter, macOS attributes the notification origin to Script Editor/osascript, thereby displaying the default Script Editor (scroll) icon. This is a default platform-specific behavior of macOS when script engines trigger user notifications, and it cannot be overriden via the command-line interface. In contrast, on Windows, notifications are triggered using the plyer library and carry the custom icon.ico file directly."
    ]
    add_custom_section(vol_iii_path, vol_iii_head, vol_iii_paras)
    
    # 2. Patch Master Document
    master_path = os.path.join(docs_dir, "System_Resource_Optimizer_Documentation.docx")
    if os.path.exists(master_path):
        add_custom_section(master_path, "Appendix G: macOS Dock Icon Caching & Notification Icon Attributes", [
            "1. macOS Launch Services Icon Cache: When compiling desktop applications with PyInstaller and Flet on macOS, the OS registers the application bundle identifier in Launch Services. Icon changes may not reflect immediately in the Dock if the cache is active. Relocating, renaming the app, or resetting the Dock process clears the cache.",
            "2. AppleScript Notification Context: Desktop alerts on macOS are routed via shell-based osascript. This executes outside the Cocoa app bundle context, causing the OS to attribute the notification banner to the Script Editor application icon rather than the custom optimizer logo. On Windows, the icon is successfully customized by passing the absolute path of the generated icon.ico file to the Windows notification tray API."
        ])

if __name__ == "__main__":
    main()
