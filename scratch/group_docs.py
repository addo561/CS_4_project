import os
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

def clean_ref_text(text):
    # Remove leading brackets and numbers, e.g. [1], [1] , [ 1 ]
    cleaned = re.sub(r'^\[\s*\d+\s*\]\s*', '', text)
    cleaned = cleaned.strip()
    return cleaned

def create_volume_cover_page(doc, volume_num, volume_title, volume_subtitle):
    print(f"Generating elegant cover page for Volume {volume_num}...")
    
    # 1. University Header
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_univ = p_univ.add_run("KWAME NKRUMAH UNIVERSITY OF SCIENCE AND TECHNOLOGY\n")
    r_univ.bold = True
    r_univ.font.size = Pt(14)
    r_univ.font.name = "Arial"
    
    p_fac = doc.add_paragraph()
    p_fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_fac = p_fac.add_run("FACULTY OF PHYSICAL AND COMPUTATIONAL SCIENCES\nDEPARTMENT OF COMPUTER SCIENCE\n")
    r_fac.bold = True
    r_fac.font.size = Pt(11)
    r_fac.font.name = "Arial"
    
    # Spacing
    for _ in range(3):
        doc.add_paragraph()
        
    # 2. Main Title (SRO Project)
    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_proj = p_proj.add_run("LIGHTWEIGHT AI-POWERED SYSTEM\nRESOURCE OPTIMIZER\n")
    r_proj.bold = True
    r_proj.font.size = Pt(18)
    r_proj.font.name = "Arial"
    r_proj.font.color.rgb = RGBColor(13, 17, 23)
    
    # Spacing
    doc.add_paragraph()
    
    # 3. Volume Title and Subtitle
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(f"VOLUME {volume_num}\n{volume_title.upper()}\n")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.name = "Arial"
    r_title.font.color.rgb = RGBColor(30, 70, 120) # Blue accent
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"{volume_subtitle}\n")
    r_sub.italic = True
    r_sub.font.size = Pt(11)
    r_sub.font.name = "Arial"
    
    # Spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # 4. Prepared By
    p_prep = doc.add_paragraph()
    p_prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_prep = p_prep.add_run("Prepared by:\n")
    r_prep.bold = True
    r_prep.font.size = Pt(11)
    r_prep.font.name = "Arial"
    
    p_students = doc.add_paragraph()
    p_students.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_std = p_students.add_run(
        "Lamptey Kwaku Abednego — 3398122\n"
        "Tugbah Lily Ama Mawuena — 3416522\n"
        "Korli Larry Addo — 3395922\n"
    )
    r_std.font.size = Pt(11)
    r_std.font.name = "Arial"
    
    # Spacing
    for _ in range(3):
        doc.add_paragraph()
        
    # 5. Footer Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run("Date: May 2026")
    r_date.bold = True
    r_date.font.size = Pt(11)
    r_date.font.name = "Arial"

def merge_volume_files(volume_num, volume_title, volume_subtitle, source_files_with_headings, output_path):
    print(f"Creating Volume {volume_num}: {volume_title}")
    if not source_files_with_headings:
        print("No files to merge.")
        return False
        
    # Start with a base document from the first source file to inherit styles, margins, page setups
    base_file = source_files_with_headings[0][0]
    merged_doc = Document(base_file)
    
    # Clear the entire body of the base document to start fresh
    print("Clearing base template body while preserving styles...")
    body = merged_doc.element.body
    for child in list(body):
        if not child.tag.endswith('sectPr'):
            body.remove(child)
            
    # Generate the Cover Page
    create_volume_cover_page(merged_doc, volume_num, volume_title, volume_subtitle)
    
    unique_ref_texts = set()
    unique_ref_list = []
    
    for i, (file_path, heading_text) in enumerate(source_files_with_headings, start=1):
        print(f"Processing and merging file {i}: {os.path.basename(file_path)}")
        
        # Add a page break before each new section
        merged_doc.add_page_break()
        
        # Open source document
        src_doc = Document(file_path)
        
        reached_content = False
        in_references = False
        
        # Scan the document to see if it contains any Heading 1 styles
        has_heading_1 = any(p.style.name == 'Heading 1' for p in src_doc.paragraphs)
        
        for child in list(src_doc.element.body):
            # Skip page-setup elements (sectPr)
            if child.tag.endswith('sectPr'):
                continue
                
            if child.tag.endswith('p'):
                p = Paragraph(child, src_doc)
                
                # Check for start of content
                if not reached_content:
                    is_start = False
                    if has_heading_1 and p.style.name == 'Heading 1':
                        is_start = True
                    elif not has_heading_1 and p.style.name == 'Normal' and p.text.strip().startswith('1. '):
                        is_start = True
                        
                    if is_start:
                        reached_content = True
                        # Substitute the heading text
                        p.style = 'Heading 1'
                        p.text = heading_text
                        print(f" -> Substituted Heading: '{heading_text}'")
                    else:
                        # Skip title block paragraphs at the beginning
                        continue
                        
                # Check for references section
                is_ref_heading = False
                if p.style.name == 'Heading 1' and ('references' in p.text.lower() or 'bibliography' in p.text.lower()):
                    is_ref_heading = True
                elif p.style.name == 'Normal' and p.text.strip().startswith(('References', 'Bibliography')):
                    is_ref_heading = True
                    
                if is_ref_heading:
                    in_references = True
                    continue
                    
                if in_references:
                    ref_text = p.text.strip()
                    if ref_text:
                        cleaned = clean_ref_text(ref_text)
                        if cleaned and cleaned not in unique_ref_texts:
                            unique_ref_texts.add(cleaned)
                            unique_ref_list.append(cleaned)
                    continue
                    
            elif child.tag.endswith('tbl'):
                if not reached_content:
                    # Skip any tables before the first heading
                    continue
                if in_references:
                    # Skip tables in references
                    continue
            
            # Insert element before sectPr to maintain valid order
            sectPr = body.sectPr
            if sectPr is not None:
                idx = body.index(sectPr)
                body.insert(idx, child)
            else:
                body.append(child)
                
    # Add consolidated references section at the very end of the volume
    if unique_ref_list:
        print(f"Merging and sorting {len(unique_ref_list)} consolidated bibliography references...")
        merged_doc.add_page_break()
        
        p_ref_head = merged_doc.add_paragraph()
        p_ref_head.style = 'Heading 1'
        p_ref_head.text = "References & Bibliography"
        
        sorted_refs = sorted(unique_ref_list, key=lambda s: s.lower())
        
        for idx, ref_text in enumerate(sorted_refs, start=1):
            p_ref = merged_doc.add_paragraph()
            p_ref.style = 'Normal'
            p_ref.paragraph_format.left_indent = Pt(18) # Hanging indent feel
            p_ref.paragraph_format.first_line_indent = Pt(-18)
            
            # Bold citation index
            r_idx = p_ref.add_run(f"[{idx}] ")
            r_idx.bold = True
            r_idx.font.name = "Arial"
            
            r_text = p_ref.add_run(ref_text)
            r_text.font.name = "Arial"
            
    print(f"Saving merged document to: {output_path}")
    merged_doc.save(output_path)
    print(f"Volume saved successfully to {os.path.basename(output_path)}!\n")
    return True

if __name__ == "__main__":
    docs_dir = "/Users/user/Desktop/Final_year/docs"
    
    # Volume I files and heading overrides
    vol_i_files = [
        (os.path.join(docs_dir, "introductions.docx"), "Section 1: Project Introduction & Final System Capabilities"),
        (os.path.join(docs_dir, "system_architecture.docx"), "Section 2: End-to-End System Architecture & Data Flow Design"),
        (os.path.join(docs_dir, "project_documentation.docx"), "Section 3: Time-Series Model and Dataset Methodology Summary")
    ]
    
    # Volume II files and heading overrides
    vol_ii_files = [
        (os.path.join(docs_dir, "dataset_collection.docx"), "Section 1: Telemetry Parameters & Dataset Collection Setup"),
        (os.path.join(docs_dir, "database_handling.docx"), "Section 2: Telemetry Data Store & Configuration Persistence Subsystem"),
        (os.path.join(docs_dir, "model_documentation.docx"), "Section 3: Gated Recurrent Unit (GRU) Model Design & Training")
    ]
    
    # Volume III files and heading overrides
    vol_iii_files = [
        (os.path.join(docs_dir, "data_pipeline.docx"), "Section 1: Real-Time Data Pipeline & Action Engine"),
        (os.path.join(docs_dir, "ui_documentation.docx"), "Section 2: Flet Design System, Thread-Safe Queues, and Canvas Charts"),
        (os.path.join(docs_dir, "mitigation_strategy.docx"), "Section 3: Cross-Platform Resource Mitigation Strategy"),
        (os.path.join(docs_dir, "presentation_deliverables.docx"), "Section 4: Academic Presentation, Live Verification & Defense Strategy")
    ]
    
    # Verify all input files exist
    all_files = [f for f, _ in vol_i_files + vol_ii_files + vol_iii_files]
    missing = [f for f in all_files if not os.path.exists(f)]
    if missing:
        print(f"Error: Missing source files: {missing}", file=sys.stderr)
        sys.exit(1)
        
    # Output paths
    output_vol_i = os.path.join(docs_dir, "Volume_I_System_Design_and_Architecture.docx")
    output_vol_ii = os.path.join(docs_dir, "Volume_II_Telemetry_and_Model_Documentation.docx")
    output_vol_iii = os.path.join(docs_dir, "Volume_III_Pipeline_UI_and_Verification.docx")
    
    # Create the volumes
    merge_volume_files("I", "Systems Engineering and Architecture Description", "Technical Specifications and System Design", vol_i_files, output_vol_i)
    merge_volume_files("II", "Telemetry Data Engineering and AI Model Development", "Data Pipelines, Storage, and GRU Training Details", vol_ii_files, output_vol_ii)
    merge_volume_files("III", "Real-Time Optimization, UI, and System Verification", "Desktop Implementation, Dashboard, and Demonstration Protocol", vol_iii_files, output_vol_iii)
    
    print("All volumes created successfully!")
